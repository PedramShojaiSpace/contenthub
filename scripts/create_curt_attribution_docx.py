from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "/home/ubuntu/lights-on-optin/docs/Curt_Interconnected_Meta_Attribution_Handoff_2026-08-17.docx"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    r.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10)
styles["Heading 1"].font.name = "Aptos Display"
styles["Heading 1"].font.size = Pt(17)
styles["Heading 1"].font.color.rgb = RGBColor(21, 75, 54)
styles["Heading 2"].font.name = "Aptos Display"
styles["Heading 2"].font.size = Pt(12)
styles["Heading 2"].font.color.rgb = RGBColor(21, 75, 54)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("THE URBAN MONK")
r.bold = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(160, 122, 46)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Interconnected Meta Attribution Handoff")
r.bold = True
r.font.size = Pt(21)
r.font.color.rgb = RGBColor(21, 75, 54)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Prepared for Curt | August 17, 2026")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(90, 90, 90)

doc.add_paragraph()

doc.add_heading("What is happening", level=1)
doc.add_paragraph(
    "Meta’s reported Lead total is not a one-to-one count of unique people who registered. "
    "For the August 10–16 reconciliation, Meta reported 4,939 lead actions while the first-party Kajabi-page registration bank contained 1,834 unique registrations. "
    "That is approximately 2.7 Meta lead actions for each unique captured registration."
)
doc.add_paragraph(
    "This gap is caused by multiple lead-related signals and platform reporting behavior—not evidence that Kajabi is losing registrations. "
    "For operating decisions, use unique first-party registrations as the lead denominator and direct Kajabi transactions as the revenue source."
)

tbl = doc.add_table(rows=1, cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Table Grid"
for c, text in zip(tbl.rows[0].cells, ["Measurement", "What It Means", "Use for Decisions?"]):
    set_cell_text(c, text, True, (255, 255, 255)); shade(c, "154B36")
rows = [
    ("Meta reported Leads", "Platform lead actions; may include multiple signals per person.", "No — diagnostic only"),
    ("First-party registrations", "Unique captured registrations in the Kajabi-page lead bank.", "Yes — lead denominator"),
    ("Direct Kajabi purchases", "Paid $67 and $199 transaction records.", "Yes — revenue source"),
]
for row in rows:
    cells = tbl.add_row().cells
    for cell, text in zip(cells, row): set_cell_text(cell, text)

doc.add_heading("Curt’s required ad-level action", level=1)
doc.add_paragraph(
    "Do not change the Meta pixel, campaign objective, ad set, creative, budget, or checkout. "
    "Only append the four campaign-identity parameters below to every active and future Agora → Interconnected destination URL."
)

doc.add_heading("Copy-ready destination URL template", level=2)
p = doc.add_paragraph()
p.style = doc.styles["Normal"]
r = p.add_run(
    "https://content.theurbanmonk.com/interconnected?utm_source=meta&utm_medium=paid_social"
    "&utm_campaign=agora_interconnected_us_2026_08&utm_content={{ad.name}}"
    "&meta_campaign_id={{campaign.id}}&meta_adset_id={{adset.id}}"
    "&meta_ad_id={{ad.id}}&meta_campaign_key=agora_interconnected_us_2026_08"
)
r.font.name = "Consolas"; r.font.size = Pt(8); r.font.color.rgb = RGBColor(35, 35, 35)

doc.add_paragraph(
    "If an ad already has the established UTM structure, keep it unchanged and append only the following suffix:",
)
p = doc.add_paragraph()
r = p.add_run("&meta_campaign_id={{campaign.id}}&meta_adset_id={{adset.id}}&meta_ad_id={{ad.id}}&meta_campaign_key=agora_interconnected_us_2026_08")
r.font.name = "Consolas"; r.font.size = Pt(8)

doc.add_heading("Required parameters", level=2)
tbl = doc.add_table(rows=1, cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Table Grid"
for c, text in zip(tbl.rows[0].cells, ["Parameter", "Value to Use", "Purpose"]):
    set_cell_text(c, text, True, (255, 255, 255)); shade(c, "154B36")
params = [
    ("meta_campaign_id", "{{campaign.id}}", "Exact Meta campaign identity"),
    ("meta_adset_id", "{{adset.id}}", "Exact Meta ad-set identity"),
    ("meta_ad_id", "{{ad.id}}", "Exact Meta ad / creative identity"),
    ("meta_campaign_key", "Stable lower-case label", "Reporting family that survives name changes"),
]
for row in params:
    cells = tbl.add_row().cells
    for cell, text in zip(cells, row): set_cell_text(cell, text)

doc.add_heading("Implementation checklist", level=1)
for item in [
    "Add the four meta_* parameters to every active Agora → Interconnected ad destination URL.",
    "Use the same meta_campaign_key for variants in the same deliberate campaign family.",
    "Use a different key for a new country, funnel, or dated test; do not use spaces.",
    "Do not alter pixels, conversion events, budgets, audience targeting, or checkout behavior as part of this URL change.",
    "After the first real opt-in from a tagged ad, notify the Content Hub operator for read-only validation.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("What will be validated after the first real opt-in", level=1)
doc.add_paragraph(
    "The first-party lead record must contain campaign ID, ad-set ID, ad ID, and campaign key alongside the existing UTM/click data. "
    "Once present, operating reports can tie unique registrations, lead → $67 conversion, $67 → $199 attachment, and earnings per lead to actual campaign and ad-set identity."
)

doc.add_heading("Current reporting rule", level=1)
doc.add_paragraph(
    "Until every active ad carries this identity, do not optimize or scale from Meta’s raw Leads column alone. "
    "Use unique first-party registrations plus direct Kajabi revenue for performance decisions; Meta lead actions remain a platform diagnostic."
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Questions or first-opt-in validation: contact the Content Hub operator before altering ad delivery settings.")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(90, 90, 90)

doc.save(OUT)
print(OUT)

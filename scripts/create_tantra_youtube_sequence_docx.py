from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "/home/ubuntu/lights-on-optin/docs/Tantra_YouTube_Sequence_and_End_Screen_Deployment_Guide.docx"
QUIZ = "https://content.theurbanmonk.com/quiz/tantra?utm_source=youtube&utm_medium=organic_video&utm_campaign=tantra_sequence"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        shade(cell, "1F4B49")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = None
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            if i == 0:
                shade(cells[i], "E7EFEA")
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    document.add_paragraph()
    return table


def add_bullet(document, text):
    document.add_paragraph(text, style="List Bullet")


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.6)
section.bottom_margin = Inches(0.6)
section.left_margin = Inches(0.65)
section.right_margin = Inches(0.65)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10)
for style_name in ["Title", "Heading 1", "Heading 2"]:
    styles[style_name].font.name = "Aptos Display"

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("TANTRA YOUTUBE SEQUENCE\nEND-SCREEN & CARD DEPLOYMENT GUIDE")
run.bold = True
run.font.size = Pt(20)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("For The Urban Monk YouTube Channel · VA Implementation Document").italic = True
doc.add_paragraph()

doc.add_heading("Objective", level=1)
doc.add_paragraph(
    "Guide viewers through seven Tantra videos as one coherent relationship-reconnection journey. "
    "The sequence earns the next view first, then offers the Tantra quiz only after the viewer has received a complete educational arc."
)

doc.add_heading("Required Playlist", level=1)
doc.add_paragraph("Create one public playlist titled: Tantra: Reconnect, Restore, and Rebuild.")
doc.add_paragraph(
    "Playlist description: A seven-part practical series from Dr. Pedram Shojai on relationship connection, desire, sensuality, health, "
    "and the small practices that help couples find their way back to one another. Begin with Video 1 and follow the sequence in order."
)

sequence = [
    ("1", "Considering Divorce? A Doctor’s Honest Take on What’s Really Happening", "Open with urgent relationship pain and create a reason to keep listening.", "Why She Stopped Wanting To"),
    ("2", "Why She Stopped Wanting To", "Validate the female experience and reframe low desire as a whole-person and relationship signal.", "Why He Stopped Wanting To"),
    ("3", "Why He Stopped Wanting To", "Make the sequence mutual rather than accusatory; explain stress and disconnection on the male side.", "The Female Orgasm: The Missing Ingredient in Western Sexuality"),
    ("4", "The Female Orgasm: The Missing Ingredient in Western Sexuality", "Introduce attentive intimacy, communication, comfort, and pleasure.", "The King and the Queen"),
    ("5", "The King and the Queen", "Expand from bedroom dynamics to the relationship and household field.", "Sex Is the Flower"),
    ("6", "Sex Is the Flower", "Connect intimacy to its roots: health, stress, sleep, inflammation, safety, and relationship systems.", "The Love Bank"),
    ("7", "The Love Bank: Why Regular Lovemaking Gives a Relationship a Longer Fuse", "Close with a sustainable practice and invite a personal next step.", "Tantra Quiz CTA"),
]

doc.add_heading("Viewer Journey Sequence", level=1)
add_table(doc, ["#", "Video", "Purpose", "End-screen destination"], sequence, [0.3, 2.0, 3.25, 1.75])

doc.add_heading("Platform Setup Rules", level=1)
add_bullet(doc, "Confirm that every video is not marked made for kids, if that accurately reflects its intended audience. End screens and cards are unavailable on made-for-kids videos.")
add_bullet(doc, "Use a 20-second end screen when the source video leaves enough open visual space. YouTube permits end screens in the final 5–20 seconds of a video.")
add_bullet(doc, "Do not put cards in the final 20 seconds; YouTube suppresses cards while end screens appear.")
add_bullet(doc, "For Videos 1–6, use three end-screen elements: the exact next video, the full playlist, and Subscribe.")
add_bullet(doc, "For Video 7, use the quiz Link end-screen element only if the channel is eligible for external links through the YouTube Partner Program. Otherwise use Playlist + Subscribe and the description/pinned comment for the quiz URL.")

doc.add_heading("Standard End-Screen Layout: Videos 1–6", level=1)
add_table(doc, ["Element", "Destination", "Placement", "Label"], [
    ("Primary video", "Exact next video in this sequence", "Largest element; left or center", "WATCH NEXT"),
    ("Playlist", "Tantra: Reconnect, Restore, and Rebuild", "Upper right", "WATCH THE FULL SERIES"),
    ("Subscribe", "The Urban Monk channel", "Lower right", "SUBSCRIBE FOR MORE"),
], [1.0, 2.6, 1.5, 1.5])

docs = [
    ("Video 1 — Considering Divorce?", "Why She Stopped Wanting To", "When desire goes quiet, there is often more to understand.", "Before you decide what this means, watch what often happens when desire goes quiet.", "considering_divorce"),
    ("Video 2 — Why She Stopped Wanting To", "Why He Stopped Wanting To", "This is not one person’s problem.", "The other side of the relationship is worth understanding too.", "why_she_stopped"),
    ("Video 3 — Why He Stopped Wanting To", "The Female Orgasm: The Missing Ingredient in Western Sexuality", "The next conversation changes the dynamic.", "The point is not blame. It is learning a more attentive conversation about intimacy.", "why_he_stopped"),
    ("Video 4 — The Female Orgasm", "The King and the Queen", "The bedroom is part of a larger relationship field.", "This works best inside a relationship culture that makes both people feel seen.", "female_orgasm"),
    ("Video 5 — The King and the Queen", "Sex Is the Flower", "Connection is supported by the roots of life together.", "Now let’s look underneath intimacy—at the roots that support or strain it.", "king_queen"),
    ("Video 6 — Sex Is the Flower", "The Love Bank", "The practical habit that gives a relationship more resilience.", "When the roots are tended, connection can become sustainable. Here is the practice that helps.", "sex_flower"),
]

doc.add_heading("Video-by-Video Deployment Sheet: Videos 1–6", level=1)
for title, next_video, teaser, prompt, utm in docs:
    doc.add_heading(title, level=2)
    doc.add_paragraph("End screen: final 20 seconds. Primary element = “%s”; secondary = full playlist; third = Subscribe." % next_video)
    doc.add_paragraph("Card: add one Video card at roughly 55–65%% of runtime, after the core educational point. Destination = “%s”." % next_video)
    doc.add_paragraph("Card teaser: %s" % teaser)
    doc.add_paragraph("End-screen prompt: %s" % prompt)
    doc.add_paragraph("Description add-on:")
    desc = doc.add_paragraph()
    desc.add_run("Watch next: ").bold = True
    desc.add_run("%s — [VA: insert the actual YouTube link]\n" % next_video)
    desc.add_run("Want a more personal starting point? Take the two-minute Tantra quiz: %s&utm_content=%s_description" % (QUIZ, utm))

doc.add_heading("Video 7 — The Love Bank: Final CTA", level=1)
doc.add_paragraph("End screen: final 20 seconds. If eligible for external links, use the quiz Link as the largest element, with Playlist upper right and Subscribe lower right.")
doc.add_paragraph("Quiz link: %s&utm_content=love_bank_end_screen" % QUIZ)
doc.add_paragraph("End-screen prompt: If you are ready for a practical next step, take the two-minute Tantra quiz. It will help you identify where to begin.")
doc.add_paragraph("If external Links are unavailable, use Playlist + Subscribe only. Put the quiz URL first in the description and pin the following comment:")
doc.add_paragraph(
    "If this video gave you a useful way to think about connection, the next step is not guesswork. Take the two-minute Tantra quiz for a practical place to start: %s&utm_content=love_bank_pinned_comment" % QUIZ,
    style="Intense Quote",
)

doc.add_heading("VA Final Checklist", level=1)
checklist = [
    "Confirm all seven videos are public or scheduled and classified appropriately for the intended audience.",
    "Create the public playlist and place videos in the exact numbered sequence.",
    "Add the standard 20-second end screen to Videos 1–6, selecting the exact next video rather than Best for viewer.",
    "Add one mid-video card to Videos 1–6 that points to the next video.",
    "Add the final CTA end screen, description link, and pinned comment to Video 7.",
    "Replace every [VA: insert …] placeholder with the finished YouTube video or playlist URL.",
    "Preview every end screen in Studio and make sure no element covers Dr. Pedram’s face, subtitles, or critical on-screen text.",
    "Test the full viewer path: Video 1 → 2 → 3 → 4 → 5 → 6 → 7 → quiz.",
    "After 7 and 30 days, review YouTube Analytics for end-screen clicks, card clicks, playlist starts, average view duration, and quiz sessions with utm_source=youtube.",
]
for item in checklist:
    add_bullet(doc, item)

doc.add_heading("Guardrails", level=1)
doc.add_paragraph(
    "Do not link a video directly to product checkout. The viewer should receive the educational sequence first, then be invited to the quiz. "
    "Do not describe the quiz as diagnostic or promise a medical outcome. Use practical, relational language such as “a personal starting point,” “identify where to begin,” and “two-minute Tantra quiz.”"
)

doc.add_heading("References", level=1)
doc.add_paragraph("[1] YouTube Help — Add end screens to videos: https://support.google.com/youtube/answer/6388789?hl=en")
doc.add_paragraph("[2] YouTube Help — Add info cards to videos: https://support.google.com/youtube/answer/6140493?hl=en")

doc.save(OUT)
print(OUT)

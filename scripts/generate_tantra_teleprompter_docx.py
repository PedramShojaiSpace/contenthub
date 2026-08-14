from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SOURCE = PROJECT_ROOT / "docs" / "tantra-video-scripts-2026-08-13.md"
OUTPUT = PROJECT_ROOT / "docs" / "tantra-sexual-health-video-scripts-teleprompter.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    tc_pr.append(shade)


def set_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Dr. Pedram Shojai  |  Tantra Video Scripts  |  Page ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def parse_scripts(markdown: str):
    pattern = re.compile(
        r"## Script (\d+) — \"([^\"]+)\"\n"
        r"\*\*Target ad sets:\*\* ([^\n]+)\n"
        r"\*\*Recommended length:\*\* ([^\n]+)\n"
        r"\*\*Blog post title:\*\* \"([^\"]+)\"\n\n---\n\n"
        r"(.*?)(?=\n---\n\n## (?:Script \d+|Blog Post Template)|\Z)",
        re.S,
    )
    scripts = []
    for match in pattern.finditer(markdown):
        scripts.append(
            {
                "number": match.group(1),
                "title": match.group(2),
                "ad_sets": match.group(3),
                "duration": match.group(4),
                "blog_title": match.group(5),
                "body": match.group(6).strip(),
            }
        )
    if len(scripts) != 7:
        raise RuntimeError(f"Expected seven scripts, found {len(scripts)}")
    return scripts


def add_script(document: Document, script: dict, is_first: bool) -> None:
    if not is_first:
        document.add_page_break()

    title = document.add_paragraph()
    title.style = document.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run(f"{script['number']}. {script['title']}")

    metadata = document.add_table(rows=2, cols=2)
    metadata.style = "Table Grid"
    metadata.autofit = False
    metadata.columns[0].width = Inches(1.35)
    metadata.columns[1].width = Inches(5.6)
    rows = [
        ("Audience", script["ad_sets"]),
        ("Run time", script["duration"]),
    ]
    for row, (label, value) in zip(metadata.rows, rows):
        set_cell_shading(row.cells[0], "264653")
        label_run = row.cells[0].paragraphs[0].add_run(label)
        label_run.font.bold = True
        label_run.font.color.rgb = RGBColor(255, 255, 255)
        value_run = row.cells[1].paragraphs[0].add_run(value)
        value_run.font.size = Pt(10.5)

    guide = document.add_paragraph()
    guide.paragraph_format.space_before = Pt(10)
    guide_run = guide.add_run("Teleprompter copy begins below. Read naturally; section cues are prompts, not spoken aloud.")
    guide_run.italic = True
    guide_run.font.size = Pt(10.5)
    guide_run.font.color.rgb = RGBColor(89, 89, 89)

    for raw_line in script["body"].splitlines():
        line = raw_line.strip()
        if not line:
            continue
        cue = re.fullmatch(r"\*\[(.+)\]\*", line)
        if cue:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(14)
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run(cue.group(1).upper())
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(38, 70, 83)
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.18
        run = paragraph.add_run(line)
        run.font.name = "Aptos"
        run.font.size = Pt(15)

    closing = document.add_paragraph()
    closing.paragraph_format.space_before = Pt(12)
    closing_run = closing.add_run("QUIZ CTA: Take the 2-minute Tantra quiz below to find your starting point.")
    closing_run.font.bold = True
    closing_run.font.size = Pt(11)
    closing_run.font.color.rgb = RGBColor(38, 70, 83)


def build_document(scripts: list[dict]) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    set_page_number(section)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(12)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = RGBColor(38, 70, 83)

    cover = document.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Inches(1.55)
    run = cover.add_run("THE URBAN MONK")
    run.font.name = "Aptos Display"
    run.font.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(38, 70, 83)

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(20)
    heading.paragraph_format.space_after = Pt(16)
    run = heading.add_run("Tantra Sexual-Health\nVideo Scripts")
    run.font.name = "Aptos Display"
    run.font.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor(38, 70, 83)

    subheading = document.add_paragraph()
    subheading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subheading.add_run("Teleprompter-ready master copy for Dr. Pedram Shojai")
    run.italic = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(89, 89, 89)

    contents = document.add_paragraph()
    contents.paragraph_format.space_before = Inches(0.85)
    contents.add_run("Included scripts\n").bold = True
    for script in scripts:
        contents.add_run(f"{script['number']}. {script['title']}\n")

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(20)
    note_run = note.add_run(
        "Use the large body copy as the spoken text. The all-caps section cues help with pacing and are not intended to be read aloud."
    )
    note_run.italic = True
    note_run.font.size = Pt(10.5)

    document.add_page_break()
    for index, script in enumerate(scripts):
        add_script(document, script, is_first=index == 0)

    document.save(OUTPUT)


def main() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    build_document(parse_scripts(markdown))
    print(OUTPUT)


if __name__ == "__main__":
    main()
